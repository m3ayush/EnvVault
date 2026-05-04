"""Generate EnvVault PRD as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


doc = Document()

# Page setup
for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────

title = doc.add_heading('EnvVault', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run('Product Requirements Document (PRD)')
sub_run.font.size = Pt(16)
sub_run.bold = True
sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_run = tagline.add_run('ML-Powered Encrypted Secrets Manager with Anomaly Detection')
tag_run.font.size = Pt(13)
tag_run.italic = True

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run('Version 1.0  |  Author: Ayushman Bhatnagar  |  April 2026')
meta_run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ── 1. EXECUTIVE SUMMARY ────────────────────────────────────────────────────

add_heading(doc, '1. Executive Summary', 1)
add_para(doc,
    'EnvVault is a full-stack, ML-powered secrets management platform designed to securely '
    'store, manage, and audit application secrets (API keys, database credentials, tokens) '
    'with intelligent anomaly detection. The system uses AES-256-GCM encryption for data '
    'at rest, role-based access control, rate limiting, and a Machine Learning service '
    'powered by Isolation Forest to detect suspicious access patterns in real-time.')

add_para(doc, 'Key Differentiators:', bold=True)
for bullet in [
    'AES-256-GCM authenticated encryption for all stored secrets',
    'Real-time anomaly detection on every access using Isolation Forest',
    'MLflow-tracked experiments with automatic model promotion (champion alias)',
    'Containerized microservices (frontend, backend, ML, MLflow) via Docker Compose',
    'Production-grade CI/CD pipeline using GitHub Actions + AWS ECR',
    'Cloud-ready infrastructure: AWS ECS Fargate + ALB via CloudFormation',
]:
    p = doc.add_paragraph(bullet, style='List Bullet')

# ── 2. PROBLEM STATEMENT ────────────────────────────────────────────────────

add_heading(doc, '2. Problem Statement', 1)
add_para(doc,
    'Modern applications require dozens of secrets (API keys, DB passwords, OAuth tokens) '
    'that are often stored in plaintext .env files, hardcoded in code, or shared via insecure '
    'channels. Even when encrypted vaults exist, organizations lack visibility into who '
    'accessed what, when, and from where — making it nearly impossible to detect '
    'credential exfiltration or insider threats in real-time.')

add_para(doc, 'EnvVault Solves:', bold=True)
for problem in [
    'Plaintext storage of sensitive credentials',
    'Lack of access auditing for security compliance',
    'No real-time detection of suspicious access patterns',
    'Difficult collaboration across developer/admin/viewer roles',
    'No automated alerting on off-hours bulk credential reads',
]:
    doc.add_paragraph(problem, style='List Bullet')

# ── 3. SYSTEM ARCHITECTURE ──────────────────────────────────────────────────

add_heading(doc, '3. System Architecture', 1)

add_heading(doc, '3.1 High-Level Architecture', 2)
add_para(doc,
    'EnvVault follows a microservices architecture with four containerized services '
    'orchestrated via Docker Compose, deployable to AWS ECS Fargate.')

# ASCII Architecture diagram in monospace
arch_diagram = """
┌────────────────────────────────────────────────────────────────┐
│                         USER (Browser)                          │
└────────────────────────────────┬───────────────────────────────┘
                                 │ HTTPS
                                 ▼
┌────────────────────────────────────────────────────────────────┐
│            React Frontend (Nginx, Port 3000)                    │
│  • Dashboard  • Audit Logs  • Secret Vault UI                   │
└────────────────────────────────┬───────────────────────────────┘
                                 │ REST API (/api/*)
                                 ▼
┌────────────────────────────────────────────────────────────────┐
│         Node.js Backend API (Express, Port 3001)                │
│  • AES-256-GCM Encryption    • RBAC (admin/dev/viewer)          │
│  • Helmet Security Headers   • Rate Limiting (100/15min)        │
│  • Audit Logging             • CORS Restriction                 │
└────────────┬───────────────────────────────────┬───────────────┘
             │                                   │
             │ POST /predict                     │ R/W
             ▼                                   ▼
┌────────────────────────────┐         ┌──────────────────────┐
│  Python ML Service         │         │  JSON File Database  │
│  Flask (Port 8000)         │         │  (db.json)           │
│  • Isolation Forest Model  │         │  • Encrypted Secrets │
│  • Real-time prediction    │         │  • Audit Logs        │
│  • Batch prediction        │         └──────────────────────┘
└────────────┬───────────────┘
             │ Track / Register
             ▼
┌────────────────────────────────────────────────────────────────┐
│            MLflow Tracking Server (Port 5050)                   │
│  • Experiment Tracking  • Model Registry  • Champion Alias      │
└────────────────────────────────────────────────────────────────┘
"""

p = doc.add_paragraph()
run = p.add_run(arch_diagram)
run.font.name = 'Courier New'
run.font.size = Pt(8)

add_heading(doc, '3.2 Component Breakdown', 2)

# Components Table
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

headers = ['Component', 'Technology', 'Port', 'Responsibility']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    set_cell_bg(hdr_cells[i], '4F81BD')

components = [
    ('Frontend', 'React + Vite + Nginx', '3000', 'User interface, secret management UI'),
    ('Backend API', 'Node.js + Express', '3001', 'Encryption, RBAC, audit logging'),
    ('ML Service', 'Python + Flask + scikit-learn', '8000', 'Anomaly detection inference'),
    ('MLflow', 'MLflow Tracking Server', '5050', 'Experiment tracking, model registry'),
]
for row_idx, comp in enumerate(components, 1):
    cells = table.rows[row_idx].cells
    for col_idx, val in enumerate(comp):
        cells[col_idx].text = val

doc.add_paragraph()

# ── 4. ML MODEL & WORKFLOW ─────────────────────────────────────────────────

add_heading(doc, '4. Machine Learning Models & Workflow', 1)

add_heading(doc, '4.1 Model: Isolation Forest', 2)
add_para(doc,
    'EnvVault uses Isolation Forest, an unsupervised ensemble algorithm specifically '
    'designed for anomaly detection. It isolates anomalies by randomly partitioning '
    'data points; outliers require fewer splits to be isolated, resulting in shorter '
    'paths in the decision trees.')

add_para(doc, 'Why Isolation Forest?', bold=True)
for reason in [
    'Unsupervised — no labeled anomaly data required',
    'Linear time complexity O(n) — scales to large datasets',
    'Effective for low anomaly rates (<5%) typical in security domains',
    'Returns continuous anomaly scores (not just binary classification)',
    'Robust to high-dimensional data',
]:
    doc.add_paragraph(reason, style='List Bullet')

add_heading(doc, '4.2 ML Pipeline', 2)
ml_diagram = """
┌─────────────────┐    ┌──────────────────┐    ┌───────────────────┐
│  Audit Logs     │───▶│  StandardScaler  │───▶│  IsolationForest  │
│  (6 features)   │    │  (Normalize)     │    │  (Detect outlier) │
└─────────────────┘    └──────────────────┘    └─────────┬─────────┘
                                                          │
                                                          ▼
                                              ┌─────────────────────┐
                                              │  Anomaly Score +    │
                                              │  Confidence Level   │
                                              └─────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(ml_diagram)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, '4.3 Feature Engineering', 2)

# Features table
feat_table = doc.add_table(rows=7, cols=3)
feat_table.style = 'Light Grid Accent 1'
hdr = feat_table.rows[0].cells
for i, h in enumerate(['Feature', 'Type', 'Normal Range / Anomaly Indicator']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

features = [
    ('hour_of_day', 'Integer 0-23', 'Normal: 8-18 (work) | Anomaly: 0-5 (late night)'),
    ('secrets_per_session', 'Integer', 'Normal: 1-8 | Anomaly: 40-60 (bulk access)'),
    ('user_role_encoded', 'Integer 0-2', '0=admin, 1=developer, 2=viewer'),
    ('ip_hash', 'Integer 1-10', 'Normal: 1-4 (office) | Anomaly: 8-10 (unknown)'),
    ('action_type_encoded', 'Integer 0-2', '0=read, 1=write, 2=delete'),
    ('day_of_week', 'Integer 0-6', 'Normal: 0-4 (weekdays) | Anomaly: 5-6 (weekends)'),
]
for row_idx, feat in enumerate(features, 1):
    cells = feat_table.rows[row_idx].cells
    for col_idx, val in enumerate(feat):
        cells[col_idx].text = val

doc.add_paragraph()

add_heading(doc, '4.4 Hyperparameter Tuning (3 Experiments)', 2)
add_para(doc, 'Three model variations are trained and tracked in MLflow, then auto-promoted using the "champion" alias:')

exp_table = doc.add_table(rows=4, cols=4)
exp_table.style = 'Light Grid Accent 1'
hdr = exp_table.rows[0].cells
for i, h in enumerate(['Run Name', 'Contamination', 'n_estimators', 'Notes']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

experiments = [
    ('baseline-v1', '0.05', '100', 'Conservative anomaly threshold'),
    ('higher-contamination-v2', '0.14', '100', 'More sensitive detection'),
    ('more-trees-v3', '0.14', '200', 'Best stability — Champion'),
]
for row_idx, exp in enumerate(experiments, 1):
    cells = exp_table.rows[row_idx].cells
    for col_idx, val in enumerate(exp):
        cells[col_idx].text = val

doc.add_paragraph()

add_heading(doc, '4.5 ML Workflow', 2)
for step in [
    'train.py — Generates synthetic audit data, trains 3 models, logs to MLflow',
    'MLflow UI — Compare runs by anomaly_rate, mean_score, hyperparameters',
    'promote_model.py — Auto-promote model with lowest anomaly_rate to "champion" alias',
    'app.py (Flask) — Loads champion model and serves /predict and /predict/batch APIs',
    'retrain.py — Triggered on data drift; trains a new model and re-promotes',
]:
    doc.add_paragraph(step, style='List Number')

# ── 5. SECURITY IMPLEMENTATION ─────────────────────────────────────────────

add_heading(doc, '5. Security Implementation', 1)

sec_table = doc.add_table(rows=8, cols=2)
sec_table.style = 'Light Grid Accent 1'
hdr = sec_table.rows[0].cells
for i, h in enumerate(['Security Layer', 'Implementation']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

security = [
    ('Encryption at Rest', 'AES-256-GCM authenticated encryption with random IV per secret'),
    ('Access Control (RBAC)', 'admin / developer / viewer roles with permission checks'),
    ('Rate Limiting', 'express-rate-limit: 100 requests per 15 min per IP'),
    ('HTTP Security Headers', 'Helmet.js (XSS, clickjacking, MIME sniffing protection)'),
    ('CORS Policy', 'Whitelist of allowed origins (no wildcard)'),
    ('Input Validation', 'Required field checks on all POST/PUT endpoints'),
    ('Audit Logging', 'Every read/write/delete logged with userId, IP, timestamp, anomaly score'),
]
for row_idx, sec in enumerate(security, 1):
    cells = sec_table.rows[row_idx].cells
    for col_idx, val in enumerate(sec):
        cells[col_idx].text = val

# ── 6. DATA FLOW ────────────────────────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '6. Application Data Flow', 1)

add_heading(doc, '6.1 Secret Creation Flow', 2)
flow1 = """
User → Frontend (React)
       │
       ▼ POST /api/secrets
    Backend (Express)
       │
       ├──▶ Validate input (keyName, value, userId)
       │
       ├──▶ Encrypt value with AES-256-GCM
       │
       ├──▶ Save to db.json
       │
       └──▶ Create audit log → POST to ML Service (/predict)
                                       │
                                       ▼
                              Isolation Forest predicts:
                              { is_anomaly, score, confidence }
                                       │
                                       ▼
                              Update audit log + alert if anomaly
"""
p = doc.add_paragraph()
run = p.add_run(flow1)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, '6.2 Anomaly Detection Flow', 2)
flow2 = """
Backend extracts features from event:
    [hour_of_day, secrets_per_session, user_role_encoded,
     ip_hash, action_type_encoded, day_of_week]
                    │
                    ▼
        POST /predict (ML Service)
                    │
                    ▼
        Isolation Forest decision_function()
                    │
                    ▼
    Score < 0  →  is_anomaly = True
    Score ≥ 0  →  is_anomaly = False
                    │
                    ▼
    Normalize to confidence % (low/medium/high)
                    │
                    ▼
    Return JSON to backend → Update audit log
"""
p = doc.add_paragraph()
run = p.add_run(flow2)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ── 7. DEVOPS & CI/CD ───────────────────────────────────────────────────────

add_heading(doc, '7. DevOps & CI/CD Pipeline', 1)

add_heading(doc, '7.1 Containerization (Docker)', 2)
add_para(doc, 'All four services run as Docker containers, orchestrated via docker-compose.yml:')
for c in [
    'frontend — React + Vite built into Nginx-served static assets',
    'backend — Node.js Alpine image with Express server',
    'ml-service — Python 3.11 with Flask + scikit-learn + MLflow',
    'mlflow — Official MLflow v2.13.0 image with persistent volume',
]:
    doc.add_paragraph(c, style='List Bullet')

add_heading(doc, '7.2 GitHub Actions CI/CD Pipeline', 2)

cicd_diagram = """
┌──────────────┐    ┌────────────────────┐    ┌────────────────────┐
│  git push    │───▶│  build-and-test    │───▶│   deploy (main)    │
│  origin main │    │                    │    │                    │
└──────────────┘    │ • npm install      │    │ • AWS Login (IAM)  │
                    │ • pip install      │    │ • ECR Login        │
                    │ • Test ML model    │    │ • Build images     │
                    │ • Build Docker x3  │    │ • Tag w/ commit    │
                    │ • Smoke test       │    │ • Push to ECR      │
                    └────────────────────┘    └────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(cicd_diagram)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para(doc, 'Pipeline triggers automatically on every push to main branch and every Pull Request.')

# ── 8. AWS CLOUD INFRASTRUCTURE ─────────────────────────────────────────────

add_heading(doc, '8. AWS Cloud Infrastructure', 1)
add_para(doc,
    'Infrastructure is defined as code via aws/cloudformation.yml — a single-command '
    'deployment that provisions the entire production stack on AWS.')

aws_table = doc.add_table(rows=8, cols=2)
aws_table.style = 'Light Grid Accent 1'
hdr = aws_table.rows[0].cells
for i, h in enumerate(['AWS Service', 'Purpose']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

aws_resources = [
    ('VPC + 2 Public Subnets', 'Network isolation across 2 availability zones'),
    ('Application Load Balancer (ALB)', 'Routes /api/* to backend, /* to frontend'),
    ('ECS Cluster (Fargate)', 'Serverless container hosting (3 services)'),
    ('ECR Repositories (x3)', 'Private Docker image registry'),
    ('S3 Bucket (versioned + AES-256)', 'ML artifacts and model storage'),
    ('IAM Roles', 'ECS execution role + task role with S3 access'),
    ('CloudWatch Log Groups', 'Centralized logs (14-day retention)'),
]
for row_idx, res in enumerate(aws_resources, 1):
    cells = aws_table.rows[row_idx].cells
    for col_idx, val in enumerate(res):
        cells[col_idx].text = val

# ── 9. API ENDPOINTS ────────────────────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '9. API Endpoints', 1)

add_heading(doc, '9.1 Backend API (Port 3001)', 2)
api_table = doc.add_table(rows=8, cols=3)
api_table.style = 'Light Grid Accent 1'
hdr = api_table.rows[0].cells
for i, h in enumerate(['Method', 'Endpoint', 'Description']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

apis = [
    ('GET', '/api/health', 'Service health check'),
    ('POST', '/api/secrets', 'Create encrypted secret'),
    ('GET', '/api/secrets', 'List secrets (RBAC enforced)'),
    ('PUT', '/api/secrets/:id/rotate', 'Rotate secret value'),
    ('DELETE', '/api/secrets/:id', 'Delete secret'),
    ('GET', '/api/audit-logs', 'Last 100 audit events'),
    ('GET', '/api/secrets/raw', 'View encrypted blobs (demo)'),
]
for row_idx, api in enumerate(apis, 1):
    cells = api_table.rows[row_idx].cells
    for col_idx, val in enumerate(api):
        cells[col_idx].text = val

doc.add_paragraph()
add_heading(doc, '9.2 ML Service API (Port 8000)', 2)

ml_api_table = doc.add_table(rows=4, cols=3)
ml_api_table.style = 'Light Grid Accent 1'
hdr = ml_api_table.rows[0].cells
for i, h in enumerate(['Method', 'Endpoint', 'Description']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

ml_apis = [
    ('GET', '/health', 'Model load status'),
    ('POST', '/predict', 'Single event anomaly prediction'),
    ('POST', '/predict/batch', 'Batch anomaly prediction'),
]
for row_idx, api in enumerate(ml_apis, 1):
    cells = ml_api_table.rows[row_idx].cells
    for col_idx, val in enumerate(api):
        cells[col_idx].text = val

# ── 10. TECH STACK ──────────────────────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '10. Technology Stack', 1)

stack_table = doc.add_table(rows=6, cols=2)
stack_table.style = 'Light Grid Accent 1'
hdr = stack_table.rows[0].cells
for i, h in enumerate(['Layer', 'Technologies']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

stack = [
    ('Frontend', 'React 18, Vite, React Router, Nginx'),
    ('Backend', 'Node.js 20, Express, Helmet, express-rate-limit, AES-256-GCM'),
    ('ML / Data Science', 'Python 3.11, scikit-learn (IsolationForest), pandas, numpy, MLflow 2.13'),
    ('DevOps', 'Docker, Docker Compose, GitHub Actions, AWS CodeBuild'),
    ('Cloud (AWS)', 'ECS Fargate, ECR, ALB, S3, IAM, CloudFormation, CloudWatch'),
]
for row_idx, s in enumerate(stack, 1):
    cells = stack_table.rows[row_idx].cells
    for col_idx, val in enumerate(s):
        cells[col_idx].text = val

# ── 11. EVALUATION CRITERIA MAPPING ─────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '11. Evaluation Criteria Mapping', 1)

eval_table = doc.add_table(rows=6, cols=3)
eval_table.style = 'Light Grid Accent 1'
hdr = eval_table.rows[0].cells
for i, h in enumerate(['Criteria', 'Max', 'Implementation']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_bg(hdr[i], '4F81BD')

evaluations = [
    ('Functional Implementation', '6', 'Full CRUD secret manager, audit logs, frontend, backend, ML — all working'),
    ('Cloud & Hyperscaler Usage', '4', 'AWS ECR (live), CloudFormation IaC, ECS Fargate, ALB, S3'),
    ('Data / ML / AI Workflow', '4', 'Isolation Forest + MLflow + 3 experiments + auto-promotion + retraining'),
    ('DevOps Practices', '3', 'Docker Compose, GitHub Actions CI/CD (live & green), AWS CodeBuild buildspec'),
    ('Security Implementation', '3', 'AES-256-GCM, RBAC, Helmet, rate limiting, CORS, audit logging'),
]
for row_idx, ev in enumerate(evaluations, 1):
    cells = eval_table.rows[row_idx].cells
    for col_idx, val in enumerate(ev):
        cells[col_idx].text = val

# ── 12. FUTURE ENHANCEMENTS ─────────────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '12. Future Enhancements', 1)
for enh in [
    'JWT-based authentication with refresh tokens',
    'OAuth2 / SSO integration (Google, GitHub)',
    'Migrate from JSON file DB to PostgreSQL or DynamoDB',
    'Add WebSocket-based real-time anomaly alerts to frontend',
    'Deploy MLflow tracking server to AWS with S3 artifact backend',
    'Implement model drift detection and automated retraining triggers',
    'Add Slack/email notifications for high-severity anomalies',
    'Multi-tenancy support with project-level isolation',
]:
    doc.add_paragraph(enh, style='List Bullet')

# ── 13. CONCLUSION ──────────────────────────────────────────────────────────

doc.add_paragraph()
add_heading(doc, '13. Conclusion', 1)
add_para(doc,
    'EnvVault demonstrates a complete production-grade software engineering project — '
    'spanning secure backend development, modern frontend architecture, machine learning '
    'integration with experiment tracking, container-based microservices, automated CI/CD, '
    'and cloud-ready infrastructure-as-code. The project showcases practical application '
    'of DevOps best practices, ML lifecycle management with MLflow, and security-first '
    'design principles suitable for real-world deployment.')

# Save
output_path = '/Users/ayushmanbhatnagar/Developer/CPPE/EnvVault_PRD.docx'
doc.save(output_path)
print(f'PRD saved to: {output_path}')
